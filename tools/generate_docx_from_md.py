#!/usr/bin/env python3
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    import subprocess, sys
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt


def parse_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    title = None
    abstract = None
    paper_parts = []

    i = 0
    n = len(lines)
    # Title
    while i < n and not title:
        line = lines[i].strip()
        if line.startswith('# '):
            title = line[2:].strip()
        i += 1

    # Find Abstract header
    i = 0
    while i < n:
        if lines[i].strip().lower().startswith('## abstract'):
            # look for blockquote line(s) after header
            j = i+1
            # gather lines until next header '##' or blank leading to next section
            abs_lines = []
            while j < n and lines[j].strip() == '':
                j += 1
            # if next line starts with '> *' it's the italicized block
            if j < n and lines[j].lstrip().startswith('>'):
                # collect consecutive blockquote lines
                while j < n and lines[j].lstrip().startswith('>'):
                    # remove leading '>' and optional space
                    txt = lines[j].lstrip()[1:].lstrip()
                    # strip wrapping asterisks if present
                    if txt.startswith('*') and txt.endswith('*'):
                        txt = txt[1:-1]
                    abs_lines.append(txt)
                    j += 1
            else:
                # collect non-empty lines until blank or next header
                while j < n and not lines[j].strip().startswith('##'):
                    if lines[j].strip() != '':
                        abs_lines.append(lines[j])
                    j += 1
            abstract = '\n'.join(abs_lines).strip()
            break
        i += 1

    # Find '## Paper' header and collect everything after as paper body
    body_start = None
    for idx, line in enumerate(lines):
        if line.strip().lower().startswith('## paper'):
            body_start = idx + 1
            break
    if body_start is None:
        # fallback: everything after title
        body_start = 0
        for idx, line in enumerate(lines):
            if line.strip().startswith('# '):
                body_start = idx + 1
                break

    # Collect body paragraphs preserving blank lines as paragraph separators
    paras = []
    cur = []
    for line in lines[body_start:]:
        if line.strip() == '':
            if cur:
                paras.append('\n'.join(cur))
                cur = []
        else:
            cur.append(line)
    if cur:
        paras.append('\n'.join(cur))

    return title or os.path.splitext(os.path.basename(path))[0], abstract, paras


def write_docx(title, abstract, paras, out_path):
    doc = Document()
    # Title
    doc.add_heading(title, level=1)

    if abstract:
        doc.add_heading('Abstract', level=2)
        p = doc.add_paragraph()
        run = p.add_run(abstract)
        run.italic = True
        # Slightly smaller font for abstract block
        try:
            run.font.size = Pt(10.5)
        except Exception:
            pass

    doc.add_heading('Paper', level=2)

    for ptext in paras:
        # If a paragraph looks like 'Author: Name' keep it on its own line and bold 'Author:'
        lines = ptext.split('\n')
        for l in lines:
            if l.strip().lower().startswith('author:'):
                run_text = l.split(':',1)
                p = doc.add_paragraph()
                r1 = p.add_run(run_text[0] + ':')
                r1.bold = True
                r1.add_break()
                r2 = p.add_run(run_text[1].lstrip())
            else:
                doc.add_paragraph(l)
        # add a blank paragraph to separate blocks
        doc.add_paragraph()

    doc.save(out_path)
    print('WROTE', out_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: generate_docx_from_md.py <input.md> [output.docx]')
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(inp)[0] + '.formatted.docx'
    title, abstract, paras = parse_markdown(inp)
    write_docx(title, abstract, paras, out)
